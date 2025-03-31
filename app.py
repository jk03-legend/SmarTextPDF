from flask import Flask, request, send_file, jsonify
from werkzeug.utils import secure_filename
import os
import docx
import requests
from pdf2docx import Converter
from flask_cors import CORS

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "/tmp"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

SAPLING_API_KEY = os.getenv("SAPLING_API_KEY")  # Store in Render environment

def extract_text_from_docx(docx_path):
    """Extracts text from a DOCX file."""
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

def proofread_text(text):
    """Proofreads text using Sapling AI API."""
    if not SAPLING_API_KEY:
        return text, [{"error": "Sapling API key is missing"}]

    url = "https://api.sapling.ai/api/v1/edits"
    headers = {"Content-Type": "application/json"}
    data = {
        "key": SAPLING_API_KEY,
        "text": text,
        "session_id": "your_session_id"
    }

    response = requests.post(url, json=data, headers=headers)
    if response.status_code != 200:
        return text, [{"error": "Failed to get response from Sapling"}]

    result = response.json()
    corrections = result.get("edits", [])

    corrected_text = text
    errors = []
    for correction in corrections:
        errors.append({
            "message": correction.get("error_type", "Grammar issue"),
            "suggestions": correction.get("replacement", []),
            "offset": correction.get("start"),
            "length": correction.get("end") - correction.get("start")
        })

    return corrected_text, errors

def save_text_to_docx(text, docx_path):
    """Saves proofread text to a new DOCX file."""
    doc = docx.Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(docx_path)

@app.route('/convert', methods=['POST'])
def convert_pdf_to_docx():
    """Handles PDF-to-DOCX conversion and proofreading."""
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    filename = secure_filename(file.filename)
    pdf_path = os.path.join(UPLOAD_FOLDER, filename)
    file.save(pdf_path)

    docx_filename = filename.rsplit('.', 1)[0] + '.docx'
    docx_path = os.path.join(OUTPUT_FOLDER, docx_filename)

    try:
        # Convert PDF to DOCX
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)  # Ensure full document conversion
        cv.close()

        if not os.path.exists(docx_path):
            return jsonify({"error": "DOCX file was not created"}), 500

        # Extract text from DOCX
        extracted_text = extract_text_from_docx(docx_path)

        # Proofread the text using Sapling
        proofread_text_content, grammar_errors = proofread_text(extracted_text)

        # Save proofread text back
