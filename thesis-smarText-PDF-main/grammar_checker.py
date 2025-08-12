import re
import os
import time
import json
import asyncio
from typing import List
from dotenv import load_dotenv

import win32com.client
from docx import Document
from docx2pdf import convert

from openai import AsyncOpenAI
from fastapi import FastAPI, HTTPException, Query
from fastapi.responses import JSONResponse
import uvicorn

# Load environment variables from .env file for API keys, etc.
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")

# Initialize OpenAI client for GPT-based grammar correction
client = AsyncOpenAI(api_key=api_key)

# ---------------------------
# PDF to Word Conversion
# ---------------------------
def convert_pdf_to_word(pdf_path, docx_path):
    """
    Convert a PDF file to a Word (.docx) file using Microsoft Word automation.
    """
    print("Converting PDF to Word using Microsoft Word...")
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(pdf_path)
    doc.SaveAs(docx_path, FileFormat=16)
    doc.Close()
    word.Quit()
    print("Conversion done.")

# ---------------------------
# Tokenizer with punctuation
# ---------------------------
def tokenize_with_punctuation(text):
    """
    Tokenizer that splits punctuation as separate tokens.
    """
    return re.findall(r"\w+|[^\w\s]", text, re.UNICODE)

# ---------------------------
# Word to PDF Conversion
# ---------------------------
def convert_word_to_pdf(updated_docx_path, final_pdf_path):
    """
    Convert a Word (.docx) file back to PDF.
    """
    print("Converting back to PDF...")
    convert(updated_docx_path, final_pdf_path)
    print(f"Final PDF saved as: {final_pdf_path}")


# ---------------------------
# GPT-based grammar correction
# ---------------------------
async def gpt_proofread(text):
    """
    Calls OpenAI GPT using a function/tool to proofread and classify changes.
    Ensures the structured JSON format and retries up to 3 times if needed.
    """

    # Normalize excessive spaces
    text = re.sub(r"[^\S\r\n]{2,}", " ", text)

    tool_schema = {
        "name": "proofread_output",
        "description": "Returns corrected content with detailed revisions for grammar, formality, and punctuation.",
        "parameters": {
            "type": "object",
            "properties": {
                "original": {"type": "string"},
                "corrected": {"type": "string"},
                "original_token": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "idx": {"type": "integer"},
                            "word": {"type": "string"}
                        },
                        "required": ["idx", "word"]
                    }
                },
                "proofread_token": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "idx": {"type": "integer"},
                            "word": {"type": "string"}
                        },
                        "required": ["idx", "word"]
                    }
                },
                "changes": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "type": {
                                "type": "string",
                                "enum": ["replaced", "corrected", "inserted", "removed"]
                            },
                            "original_idx": {"type": ["integer", "null"]},
                            "proofread_idx": {"type": ["integer", "null"]},
                            "original_word": {"type": "string"},
                            "proofread_word": {"type": "string"},
                            "suggestion": {
                            "type": "array",
                            "items": {
                                "type": "string",
                                "pattern": "^[^\\s]+$"
                            },
                            "maxItems": 3
                            }
                        },
                        "required": [
                            "type", "original_idx", "proofread_idx",
                            "original_word", "proofread_word", "suggestion"
                        ]
                    }
                }
            },
            "required": ["original", "corrected", "original_token", "proofread_token", "changes"]
        }
    }


    system_msg = (
        "You are a professional AI assistant integrated into a proofreading system called SmarText PDF. \n"
        "You will receive raw text extracted from PDFs and must correct grammar, punctuation, and spelling errors. \n"
        "Improve the tone to be more formal and ensure the text is clear and coherent, while keeping the original meaning intact. "
        "You may receive content in various languages. Proofread in the language the text is written in — be bilingual and flexible. \n"
        "Do not translate unless explicitly asked. Your goal is to refine the original language, not convert it. \n\n"

        "Do not modify names of people, places, companies, or titles — even if they appear to be misspelled, uncapitalized, or stylized. "
        "Preserve them exactly as they appear in the original text, unless the correction is absolutely unambiguous (e.g., fixing 'jonh' to 'John' when the context clearly indicates a typo). "
        "Avoid guessing or rewriting names, especially rare, creative, or user-generated ones like handles or nicknames. \n\n"

        "Examples:\n"
        "- Leave names like 'john smith', 'Gooogle', 'mcdonalds', or 'Elonn' unchanged.\n"
        "- Do not correct 'Jhon' to 'John' unless it's clearly a typo.\n"
        "- Do not change 'may' to 'might' if 'May' is a name or month.\n"
        "- Do not alter stylized names like 'iPhone', 'eBay', or usernames like '@daniel42'.\n\n"

        "Your output must strictly use the function proofread_output with all required fields: "
        "original, corrected, original_token, proofread_token, and changes.\n\n"

        "Each token must be processed as a word or punctuation mark — no grouping or skipping. "
        "If punctuation changes (e.g., '.', ',', '?', '!') or is added/removed, you must include it in changes. "
        "If a token is part of a numbered heading or subheading (e.g., '1 ', '1.', '1.1.', 'I.', 'A)'), retain the numbering or Roman numeral as-is. Do not modify or remove it. "
        "Maintain accurate idx for both original_token and proofread_token. \n\n"

        "For each change:\n"
        "- Use corrected for fixing grammar or punctuation with minor word adjustments.\n"
        "- Use replaced when words or phrases are changed into different terms (and optionally include up to 3 suggestions).\n"
        "- Use inserted if new words were added.\n"
        "- Use removed if unnecessary words or punctuation were deleted.\n\n"

        "Every change must include original_word, proofread_word, and suggestion even if suggestions only include the accepted version. "
        "Ensure the suggestion list is meaningful — if only one correction exists, include it alone. If multiple rewrites are possible, list alternatives.\n\n"

        "Do not skip punctuation or formatting changes. The goal is accurate proofreading that can be traced token-by-token and visually rendered with detailed changes."
    )



    def ensure_structure(result, original_text):
        return {
            "original": result.get("original", original_text),
            "corrected": result.get("corrected", original_text),
            "original_token": result.get("original_token", []),
            "proofread_token": result.get("proofread_token", []),
            "changes": result.get("changes", [])
        }

    for attempt in range(1, 4):
        try:
            response = await client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_msg},
                    {"role": "user", "content": f"Proofread this sentence:\n{text}"}
                ],
                tools=[{"type": "function", "function": tool_schema}],
                tool_choice="auto",
                temperature=0.3,
            )

            function_call = response.choices[0].message.tool_calls[0].function
            result = json.loads(function_call.arguments)
            result = ensure_structure(result, text)

            # Check completeness
            if result["original_token"] and result["proofread_token"]:
                return result
            else:
                print(f"[WARNING] Attempt {attempt}: Missing token data.")
        except Exception as e:
            print(f"[ERROR] Attempt {attempt}: Tool call failed.\nInput: {repr(text[:200])}")
            print("Error:", e)

    # Fallback return
    print("[ERROR] All 3 attempts failed. Returning fallback structure.")
    return {
        "original": text,
        "corrected": text,
        "original_token": [],
        "proofread_token": [],
        "changes": []
    }
# ---------------------------
# Async wrapper for GPT proofreading
# ---------------------------
async def async_gpt_proofread(paragraph_id, text):
    """
    Async wrapper to call GPT proofreading for a paragraph.
    Returns paragraph_id, original text, and GPT response.
    """
    try:
        gpt_response = await gpt_proofread(text)
        return paragraph_id, text, gpt_response
    except Exception as e:
        print(f"Error processing paragraph {paragraph_id}: {e}")
        return paragraph_id, text, {"corrected": text}
    
# ---------------------------
# Main async grammar correction for all paragraphs
# ---------------------------
async def correct_paragraphs_async(
    docx_path, updated_docx_path, json_output_path, pdf_id="example_pdf_001"
):
    """
    Proofreads all paragraphs in a Word document asynchronously using GPT,
    updates the document, and writes a JSON report.
    Returns the number of improved paragraphs.
    """
    doc = Document(docx_path)
    data = {"pdf_id": pdf_id, "paragraphs": []}
    total_word_changes = 0

    tasks = []
    paragraph_map = []

    # Step 1: prepare tasks for each paragraph
    for idx, paragraph in enumerate(doc.paragraphs):
        original_text = paragraph.text
        if not original_text.strip():
            continue
        para_id = idx + 1
        paragraph_map.append((idx, para_id, paragraph))
        
        # DEBUGGING: Log the paragraph that's about to be proofread
        print(f"Proofreading Paragraph {para_id}: {original_text}")
        tasks.append(async_gpt_proofread(para_id, original_text))

    # Step 2: execute all tasks concurrently
    results = await asyncio.gather(*tasks)
    
    # print(results)

      # Step 3: update paragraphs and prepare report
    for (idx, para_id, paragraph), (returned_para_id, original_text, gpt_response) in zip(paragraph_map, results):
        corrected_text = gpt_response.get("corrected", original_text)

        data["paragraphs"].append(
            {
                "paragraph_id": para_id,
                "original": gpt_response.get("original"),
                "proofread": gpt_response.get("corrected"),
                "original_token": gpt_response.get("original_token", []),
                "proofread_token": gpt_response.get("proofread_token", []),
                "original_text": [
                    {
                        "index": ch.get("original_idx"),
                        "word": ch.get("original_word"),
                        "type": "error",
                    }
                    for ch in gpt_response.get("changes", [])
                    if ch.get("type") != "inserted"
                    and ch.get("original_idx") is not None
                ],
                "revised_text": [
                    {
                        "index": ch.get("proofread_idx"),
                        "word": ch.get("proofread_word"),
                        "type": ch.get("type"),
                        "suggestions": ch.get("suggestion", [ch.get("proofread_word")]),
                    }
                    for ch in gpt_response.get("changes", [])
                    if ch.get("proofread_idx") is not None
                ],
            }
        )

        # Update Word paragraph while preserving formatting
        if paragraph.runs:
            ref_run = paragraph.runs[0]
            paragraph.clear()
            new_run = paragraph.add_run(corrected_text)
            new_run.font.name = ref_run.font.name
            new_run.bold = ref_run.bold
            new_run.italic = ref_run.italic
            new_run.underline = ref_run.underline
            new_run.font.size = ref_run.font.size
            if ref_run.font.color and ref_run.font.color.rgb:
                new_run.font.color.rgb = ref_run.font.color.rgb
        else:
            paragraph.text = corrected_text

        # Count word changes for this paragraph and accumulate
        word_changes_count = sum(
            1 for ch in gpt_response.get("changes", [])
            if ch.get("type") in ["inserted", "changed", "replaced", "corrected"]
        )
        total_word_changes += word_changes_count

    doc.save(updated_docx_path)
    print("Document updated with grammar corrections.")

    with open(json_output_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    print(f"Proofread JSON saved to {json_output_path}")

    return total_word_changes

# ---------------------------
# Update selected paragraphs in DOCX from JSON
# ---------------------------
def update_changes_on_pdf(
    final_pdf_path, updated_docx_path, json_output_path, paragraph_id
):
    """
    Updates the paragraphs in a Word document based on proofread data from a JSON file.
    Only updates paragraphs whose IDs are in paragraph_id.
    Returns the number of paragraphs updated.
    """
    # Load proofreading results from JSON
    with open(json_output_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    paragraphs_data = data.get("paragraphs", [])
    paragraph_id_set = set(paragraph_id)

    # Load the existing DOCX file
    doc = Document(updated_docx_path)
    updated_count = 0

    for para in paragraphs_data:
        pid = str(para.get("paragraph_id"))
        if pid in paragraph_id_set:
            print(f"Now Processing Paragraph ID : {pid}")
            proofread_text = para.get("proofread")
            para_index = int(pid) - 1  # Adjust for 0-based indexing

            if 0 <= para_index < len(doc.paragraphs):
                paragraph = doc.paragraphs[para_index]
                if paragraph.runs:
                    ref_run = paragraph.runs[0]
                    paragraph.clear()
                    new_run = paragraph.add_run(proofread_text)
                    new_run.font.name = ref_run.font.name
                    new_run.bold = ref_run.bold
                    new_run.italic = ref_run.italic
                    new_run.underline = ref_run.underline
                    new_run.font.size = ref_run.font.size
                    if ref_run.font.color and ref_run.font.color.rgb:
                        new_run.font.color.rgb = ref_run.font.color.rgb
                else:
                    paragraph.text = proofread_text
                updated_count += 1

    # Save changes to DOCX
    doc.save(updated_docx_path)

    return updated_count

# ---------------------------
# FastAPI app and endpoint
# ---------------------------
app = FastAPI()

@app.get("/api/grammar-check")
async def grammar_check(
    mode: int = Query(...), file_code: str = Query(...), paragraph_id: str = Query(...)
):
    """
    Main API endpoint for grammar checking and PDF processing.
    - mode="0": Full process (PDF→Word→Proofread→PDF)
    - mode="1": Update only selected paragraphs (using paragraph_id)
    Returns output filenames, total improvements, and elapsed time.
    """
    print(
        f"Received request: mode={mode}, file_code={file_code}, paragraph_id={paragraph_id}"
    )

    # Clean input like "[1,2,3]" or "1,2,3"
    cleaned = re.sub(r"[\[\]\s]", "", paragraph_id)

    start_time = time.time()

    pdf_path = os.path.abspath(f"original_pdfs/{file_code}.pdf")
    docx_path = os.path.abspath(f"parsing_words/{file_code}_temp.docx")
    updated_docx_path = os.path.abspath(f"parsing_words/{file_code}_updated.docx")
    final_pdf_path = os.path.abspath(f"processed_pdfs/{file_code}.pdf")
    json_output_path = os.path.abspath(f"jsons/{file_code}.json")

    if mode == 0:
        # Full process: convert, proofread, and save all
        convert_pdf_to_word(pdf_path, docx_path)
        total_improvements = await correct_paragraphs_async(
            docx_path, updated_docx_path, json_output_path
        )
        convert_word_to_pdf(updated_docx_path, final_pdf_path)
    else:
        # Only update selected paragraphs
        convert_pdf_to_word(final_pdf_path, docx_path)
        total_improvements = update_changes_on_pdf(
            final_pdf_path, updated_docx_path, json_output_path, paragraph_id
        )
        convert_word_to_pdf(updated_docx_path, final_pdf_path)

    elapsed = time.time() - start_time
    print(f"Total processing time: {elapsed:.2f} seconds")
    print(f"Total Improvements Found: {total_improvements}")

    return {
        "json_filename": os.path.basename(json_output_path),
        "final_pdf_filename": os.path.basename(final_pdf_path),
        "total_improvements": total_improvements,
        "elapsed_time_seconds": round(elapsed, 2),
    }

if __name__ == "__main__":
    # Run the FastAPI app with Uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5000)